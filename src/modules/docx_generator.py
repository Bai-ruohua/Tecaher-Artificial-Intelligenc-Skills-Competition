# -*- coding: utf-8 -*-
"""
DOCX 文档生成模块 v4.7
支持 3 种模板：教案（按章节）/ 授课计划（整门课）/ 课标（整门课）
所有输出保持原模板格式，仅替换内容文本。
AI 失败时自动降级为基于章节数据的模板内容。
"""
import os
import re
import json
import shutil
from docx.shared import Pt
from config import Config
from database import get_course_by_id, get_course_chapter, get_course_chapters
from deepseek_client import deepseek

TEMPLATE_DIR = os.path.join(Config.BASE_DIR, "data", "templates")


# ========== 工具函数 ==========

def _template_path(name):
    """获取模板绝对路径"""
    p = os.path.join(TEMPLATE_DIR, name)
    if not os.path.exists(p):
        src_map = {
            "教案.docx": r"D:\xwechat_files\wxid_z87ix56ty7je22_de07\msg\attach\4b570798299f791ce27783b13a4311d8\2026-07\Rec\81431309cf4af7ff\F\1\物联网设备装调与维护-教案.docx",
            "授课计划.docx": r"D:\xwechat_files\wxid_z87ix56ty7je22_de07\msg\attach\4b570798299f791ce27783b13a4311d8\2026-07\Rec\81431309cf4af7ff\F\0\物联网设备装调与维护-授课计划.docx",
            "课标.docx": r"D:\xwechat_files\wxid_z87ix56ty7je22_de07\msg\attach\4b570798299f791ce27783b13a4311d8\2026-07\Rec\81431309cf4af7ff\F\2\物联网设备装调与维护-课标.docx",
        }
        src = src_map.get(name)
        if src and os.path.exists(src):
            os.makedirs(TEMPLATE_DIR, exist_ok=True)
            shutil.copy2(src, p)
    return p if os.path.exists(p) else None


def _course_name(course_id):
    try:
        c = get_course_by_id(course_id)
        return c["course_name"] if c else "本课程"
    except Exception:
        return "本课程"


def _course_textbook(course_id):
    """获取教材信息"""
    try:
        c = get_course_by_id(course_id)
        tb = c.get("textbook", "") if c else ""
        return tb if tb else ""
    except Exception:
        return ""


def _fill_cell(cell, text, size=10):
    """替换单元格内容，保留单元格格式"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    return run


def _fill_cell_by_keyword(row, keyword, value, size=10):
    """在某行中查找包含 keyword 的单元格，替换内容"""
    from docx.shared import Pt
    for cell in row.cells:
        if keyword in cell.text:
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(size)
            return True
    return False


def _save_docx(doc, course_id, suffix):
    """保存 DOCX 到 uploads 目录，返回下载 URL"""
    from docx.shared import Pt
    cname = re.sub(r'[\\/:*?"<>|]', '_', _course_name(course_id))
    out_dir = os.path.join(Config.UPLOAD_DIR, str(course_id))
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{cname}-{suffix}.docx")
    if os.path.exists(out_path):
        n = 1
        while os.path.exists(os.path.join(out_dir, f"{cname}-{suffix}_{n}.docx")):
            n += 1
        out_path = os.path.join(out_dir, f"{cname}-{suffix}_{n}.docx")
    doc.save(out_path)
    return f"/api/course/{course_id}/download/{os.path.basename(out_path)}"


# ========== JSON 序列化辅助 ==========

def _ensure_list(val):
    """确保值是一个列表（兼容字符串存储的 JSON 列表）"""
    if isinstance(val, list):
        return val
    if isinstance(val, str):
        try:
            return json.loads(val)
        except Exception:
            return [val] if val else []
    return []


# ========== AI 内容生成（带降级）==========

def _call_ai(system, prompt, max_tokens=4096):
    """调用 DeepSeek AI，失败时返回 None"""
    try:
        result = deepseek.chat(system, prompt, max_tokens=max_tokens)
        result = result.strip()
        if result.startswith("```"):
            result = result.split("\n", 1)[1] if "\n" in result else result[3:]
            if result.endswith("```"):
                result = result[:-3]
            result = result.strip()
        return result
    except Exception:
        return None


def _parse_json_ai_response(raw):
    """解析 AI 返回的 JSON 字符串"""
    if not raw:
        return None
    try:
        return json.loads(raw)
    except Exception:
        return None


def _ai_or_fallback(ai_func, fallback_func):
    """调用 AI，失败时使用降级函数"""
    result = ai_func()
    if result is not None:
        return result
    return fallback_func()


# ========== 课程信息获取 ==========

def _get_course_info(course_id):
    """获取课程名称、教材、章节列表"""
    cname = _course_name(course_id)
    textbook = _course_textbook(course_id)
    chapters = get_course_chapters(course_id) or []
    return cname, textbook, chapters


# ========== 降级内容生成（无需 AI）==========

def _fallback_lesson_plan(ch, cname, textbook=""):
    """基于章节数据生成教案内容，无需 AI"""
    objectives = _ensure_list(ch.get('objectives', []))
    key_points = _ensure_list(ch.get('key_points', []))
    difficulties = _ensure_list(ch.get('difficulties', []))

    kp_text = "；".join(key_points) if key_points else ch.get('title', '')
    ob_text = "；".join(objectives) if objectives else f"掌握{ch.get('title', '')}的基本概念和操作方法"
    df_text = "；".join(difficulties) if difficulties else kp_text

    tb_suffix = f"参考教材《{textbook}》" if textbook else ""

    return {
        "teaching_content_analysis": f"本章节「{ch['title']}」是{cname}课程的核心内容，主要讲解{kp_text}。内容组织遵循由浅入深、理论联系实际的原则。{tb_suffix}",
        "student_analysis": "学生已具备基本的计算机操作能力，对本课程有较强学习兴趣，但部分学生基础薄弱，需加强实践环节的引导。",
        "knowledge_objective": ob_text,
        "ability_objective": f"能够运用{ch['title']}相关知识分析和解决实际问题",
        "literacy_objective": "培养严谨求实的科学态度和团队协作的工程素养",
        "key_points": kp_text,
        "difficulties": df_text,
        "teaching_strategy": "采用案例教学+任务驱动模式，结合多媒体课件和在线编程实践平台，以学生为中心、教师引导为辅。",
        "before_class": f"预习教材中关于{ch['title']}的相关内容，了解基本概念。{tb_suffix}",
        "intro": f"通过一个与{ch['title']}相关的实际案例引入本课内容，激发学生学习兴趣。",
        "announce": "明确本课的学习目标和重点难点。",
        "test": "通过课堂提问和小测验检测学生预习效果。",
        "teach": f"系统讲解{ch['title']}的核心知识点：{kp_text}。结合案例演示操作步骤。",
        "practice": f"布置课堂练习，让学生动手实践{ch['title']}的相关操作。教师巡回指导。",
        "summary": f"总结本章节{ch['title']}的核心知识点，强调重点和易错点。",
        "homework": f"完成课后习题，巩固{ch['title']}知识。{tb_suffix}",
    }


def _ai_generate_lesson_plan(chapter_id, course_id, extra=""):
    """生成教案内容（JSON），失败时返回 None"""
    ch = get_course_chapter(course_id, chapter_id)
    if not ch:
        return None
    cname = _course_name(course_id)
    textbook = _course_textbook(course_id)
    tb_hint = f"\n参考教材：{textbook}" if textbook else ""
    prompt = f"""你是一位高职课程设计专家。请为课程「{cname}」的章节「{ch['title']}」生成教案内容。

章节信息：
- 教学目标：{'；'.join(_ensure_list(ch.get('objectives', [])))}
- 重点：{'；'.join(_ensure_list(ch.get('key_points', [])))}
- 难点：{'；'.join(_ensure_list(ch.get('difficulties', [])))}
- 级别：{ch.get('level', '')}
{tb_hint}
{('特殊要求：'+extra) if extra else ''}

请按 JSON 格式输出：
{{
  "teaching_content_analysis": "教学内容分析（150字）",
  "student_analysis": "学情分析（150字）",
  "knowledge_objective": "知识目标",
  "ability_objective": "能力目标",
  "literacy_objective": "素养目标",
  "key_points": "教学重点",
  "difficulties": "教学难点",
  "teaching_strategy": "教学策略",
  "before_class": "课前预习任务",
  "intro": "导入学习情境描述",
  "announce": "宣布教学目标",
  "test": "目标检测方式",
  "teach": "教师讲授内容",
  "practice": "学生巩固环节",
  "summary": "总结反思",
  "homework": "课后作业"
}}"""
    raw = _call_ai("你是一个严谨的JSON输出助手。", prompt, max_tokens=4096)
    return _parse_json_ai_response(raw)


def _fallback_teaching_plan(course_id):
    """基于章节数据生成授课计划内容，无需 AI"""
    cname, textbook, chapters = _get_course_info(course_id)
    n = len(chapters)
    total_hours = max(n * 4, 36)
    schedule = []
    for i, ch in enumerate(chapters):
        week = i + 1
        module = ch.get('module', '') or f"模块{week}"
        schedule.append({
            "week": week,
            "class_num": f"{week*4-3}-{week*4}",
            "module": module,
            "task": ch.get('chapter_id', ''),
            "content": ch.get('title', ''),
            "requirement": f"掌握{'；'.join(_ensure_list(ch.get('key_points', []))[:3])}" if ch.get('key_points') else "掌握基本概念和操作",
            "ideology": "培养严谨的科学态度和工匠精神",
            "theory_hours": 2,
            "practice_hours": 2,
            "method": "讲授法、讨论法、直观演示法、练习法",
        })
    return {
        "course_name": cname,
        "total_hours": str(total_hours),
        "theory_hours": str(total_hours // 3),
        "practice_hours": str(total_hours * 2 // 3),
        "semester": "2026学年第2学期",
        "textbook": textbook or "相关教材",
        "teaching_objectives": f"通过本课程学习，使学生掌握{cname}的基本概念、原理和方法，具备运用相关知识解决实际问题的能力。",
        "teaching_content": f"本课程共{n}个章节，涵盖{cname}的核心知识体系，采用项目化教学方式组织教学内容。",
        "teaching_method": "项目教学法、案例教学法、任务驱动法相结合",
        "assessment": "过程性考核（平时成绩40%）+ 终结性考核（期末考试60%）",
        "schedule": schedule,
    }


def _ai_generate_teaching_plan(course_id):
    """生成授课计划内容，失败时返回 None"""
    cname, textbook, chapters = _get_course_info(course_id)
    ch_list = [f"{c['chapter_id']} {c['title']}" for c in chapters[:20]]
    ch_text = "\n".join(f"- {t}" for t in ch_list)
    tb_hint = f"\n参考教材：{textbook}" if textbook else ""

    prompt = f"""你是一位高职课程负责人。请为课程「{cname}」生成授课计划内容。

课程章节：
{ch_text}
{tb_hint}

请按 JSON 格式输出：
{{
  "course_name": "{cname}",
  "total_hours": "72",
  "theory_hours": "24",
  "practice_hours": "48",
  "semester": "2026学年第2学期",
  "textbook": "相关教材",
  "teaching_objectives": "教学目标（200字，以能力描述）",
  "teaching_content": "教学内容（150字）",
  "teaching_method": "教学方法建议",
  "assessment": "考核与评价方式",
  "schedule": [
    {{"week":1,"class_num":"1-4","module":"项目一","task":"1.1-1.2","content":"教学内容描述","requirement":"教学要求描述","ideology":"思政主题","theory_hours":2,"practice_hours":2,"method":"讲授法、讨论法"}},
    {{"week":2,"class_num":"5-8","module":"项目二","task":"2.1-2.2","content":"...","requirement":"...","ideology":"...","theory_hours":2,"practice_hours":2,"method":"..."}}
  ]
}}
要求：schedule 数组根据实际章节数量生成，每章4-8学时。week从1递增。"""
    raw = _call_ai("你是一个严谨的JSON输出助手。", prompt, max_tokens=4096)
    return _parse_json_ai_response(raw)


def _fallback_curriculum_standard(course_id):
    """基于章节数据生成课程标准内容，无需 AI"""
    cname, textbook, chapters = _get_course_info(course_id)
    n = len(chapters)
    total_hours = max(n * 4, 36)

    modules = []
    for i, ch in enumerate(chapters[:7]):
        modules.append({
            "name": ch.get('module', '') or f"项目{i+1}",
            "hours": 4 if i < 7 else total_hours - 28,
            "objectives": f"掌握{'；'.join(_ensure_list(ch.get('key_points', []))[:2])}" if ch.get('key_points') else f"掌握{ch['title']}",
            "content": ch['title'],
        })

    return {
        "course_name": cname,
        "major": "相关专业",
        "course_type": "专业核心课",
        "grade": "大二年级",
        "total_hours": total_hours,
        "credits": 4,
        "knowledge_objectives": [f"掌握{cname}的基本概念和原理", f"熟悉{len(chapters)}个章节的核心知识点", "具备独立分析和解决实际问题的能力"],
        "ability_objectives": ["能够运用课程知识完成项目开发或数据分析任务", "能够阅读和理解相关技术文档", "具备自主学习和技术创新的能力"],
        "literacy_objectives": ["培养严谨的科学态度和工匠精神", "树立正确的职业道德观", "培养团队协作和沟通能力"],
        "prerequisites": [{"name": "计算机应用基础", "support": "具备基本计算机操作能力"}, {"name": "高等数学（可选）", "support": "具备基本逻辑思维能力"}],
        "modules": modules,
        "assessment": "过程性考核（40%）+ 终结性考核（60%）",
        "assessment_items": [
            {"type": "过程性考核", "project": "课堂表现", "content": "出勤、课堂互动、小组讨论", "method": "教师评价", "weight": "10%"},
            {"type": "过程性考核", "project": "平时作业", "content": "课后习题、实践报告", "method": "教师批改", "weight": "20%"},
            {"type": "过程性考核", "project": "阶段测试", "content": "期中考试或阶段性测验", "method": "闭卷笔试", "weight": "10%"},
            {"type": "终结性考核", "project": "期末考试", "content": "课程全部内容", "method": "闭卷笔试（60%）+ 上机操作（40%）", "weight": "60%"},
        ],
        "teacher_requirements": "具备相关专业背景，有企业实践经历者优先",
        "equipment": "计算机实训室（每人一台）、多媒体教学设备、在线教学平台",
    }


def _ai_generate_curriculum_standard(course_id):
    """生成课程标准内容，失败时返回 None"""
    cname, textbook, chapters = _get_course_info(course_id)
    ch_text = "\n".join(f"- {c['chapter_id']} {c['title']}（{c.get('level','')}）" for c in chapters[:20])
    tb_hint = f"\n参考教材：{textbook}" if textbook else ""

    prompt = f"""你是一位高职课程专家。请为课程「{cname}」生成课程标准内容。

课程章节：
{ch_text}
{tb_hint}

请按 JSON 格式输出：
{{
  "course_name": "{cname}",
  "major": "相关专业",
  "course_type": "专业核心课",
  "grade": "大二年级",
  "total_hours": 72,
  "credits": 4,
  "knowledge_objectives": ["知识目标1","知识目标2","知识目标3"],
  "ability_objectives": ["能力目标1","能力目标2","能力目标3"],
  "literacy_objectives": ["素质目标1","素质目标2","素质目标3"],
  "prerequisites": [{{"name":"先修课程1","support":"支撑能力说明"}}],
  "modules": [{{"name":"模块1","hours":8,"objectives":"目标","content":"内容"}}],
  "assessment": "考核方式说明",
  "assessment_items": [{{"type":"过程性考核","project":"考核项目","content":"考核内容","method":"考核方式","weight":"权重"}}],
  "teacher_requirements": "教师团队要求",
  "equipment": "实训室硬件配置要求"
}}"""
    raw = _call_ai("你是一个严谨的JSON输出助手。", prompt, max_tokens=4096)
    return _parse_json_ai_response(raw)


def _fallback_course_lesson_plan(course_id):
    """基于章节数据生成整门课教案内容，无需 AI"""
    cname, textbook, chapters = _get_course_info(course_id)
    result = []
    for ch in chapters:
        result.append(_fallback_lesson_plan(ch, cname, textbook))
    return result


def _ai_generate_course_lesson_plan(course_id):
    """生成整门课所有章节的教案内容，失败时返回 None"""
    cname, textbook, chapters = _get_course_info(course_id)
    if not chapters:
        return None

    ch_summary = "\n".join(
        f"- {c['chapter_id']} {c['title']}（重点：{'；'.join(_ensure_list(c.get('key_points', []))[:2])}）"
        for c in chapters
    )
    tb_hint = f"\n参考教材：{textbook}" if textbook else ""

    prompt = f"""你是一位高职课程设计专家。请为课程「{cname}」生成整门课的教案内容。

课程包含以下章节：
{ch_summary}
{tb_hint}

请按 JSON 格式输出，每个章节一个对象：
[
  {{
    "chapter_id": "ch01",
    "teaching_content_analysis": "教学内容分析（80字）",
    "student_analysis": "学情分析（80字）",
    "knowledge_objective": "知识目标",
    "ability_objective": "能力目标",
    "literacy_objective": "素养目标",
    "key_points": "教学重点",
    "difficulties": "教学难点",
    "teaching_strategy": "教学策略（50字）",
    "before_class": "课前预习",
    "intro": "导入情境",
    "announce": "教学目标宣布",
    "test": "目标检测",
    "teach": "教师讲授内容",
    "practice": "学生巩固环节",
    "summary": "总结反思",
    "homework": "课后作业"
  }}
]"""

    raw = _call_ai("你是一个严谨的JSON输出助手。", prompt, max_tokens=8192)
    if raw:
        data = _parse_json_ai_response(raw)
        if isinstance(data, list):
            return data
    return None


# ========== 教案生成（单章节）==========

def generate_lesson_plan_docx(chapter_id, course_id, extra=""):
    """基于教案.docx模板，按章节生成教案"""
    from docx import Document
    from docx.shared import Pt

    tmpl = _template_path("教案.docx")
    if not tmpl:
        return {"success": False, "error": "教案模板不存在"}

    ch = get_course_chapter(course_id, chapter_id)
    if not ch:
        return {"success": False, "error": "章节不存在"}

    doc = Document(tmpl)
    cname = _course_name(course_id)
    textbook = _course_textbook(course_id)

    # AI 生成（带降级）
    content = _ai_generate_lesson_plan(chapter_id, course_id, extra)
    if not content:
        content = _fallback_lesson_plan(ch, cname, textbook)

    tables = doc.tables

    # Table0: 教案主表（25行×8列）
    if len(tables) > 0:
        t0 = tables[0]
        rows = t0.rows
        if len(rows) > 0:
            for cell in rows[0].cells:
                t = cell.text.strip()
                if "项目" in t or "课题" in t:
                    _fill_cell(cell, f"项目 {ch['title']}", 10)
                elif "学时" in t:
                    _fill_cell(cell, "4", 10)
        if len(rows) > 1:
            _fill_cell_by_keyword(rows[1], "学情分析", content.get("student_analysis", ""), 9)
        if len(rows) > 2:
            _fill_cell_by_keyword(rows[2], "思想政治", content.get("literacy_objective", ""), 9)
        if len(rows) > 3:
            _fill_cell_by_keyword(rows[3], "知识目标", content.get("knowledge_objective", ""), 9)
        if len(rows) > 4:
            _fill_cell_by_keyword(rows[4], "能力目标", content.get("ability_objective", ""), 9)
        if len(rows) > 5:
            _fill_cell_by_keyword(rows[5], "思政目标", content.get("literacy_objective", ""), 9)
        if len(rows) > 6:
            _fill_cell_by_keyword(rows[6], "教学重点", content.get("key_points", ""), 9)
        if len(rows) > 7:
            _fill_cell_by_keyword(rows[7], "教学难点", content.get("difficulties", ""), 9)
        if len(rows) > 8:
            _fill_cell_by_keyword(rows[8], "考核要点", content.get("knowledge_objective", ""), 9)
        if len(rows) > 9:
            _fill_cell_by_keyword(rows[9], "教学方法", content.get("teaching_strategy", ""), 9)
        if len(rows) > 10:
            _fill_cell_by_keyword(rows[10], "课程类型", "实践+理论（4学时）", 9)
        if len(rows) > 11:
            pass  # 教学资源 - 保留模板原内容

    # 教学流程填充（行13以后）
    if len(tables) > 0:
        rows = tables[0].rows
        flow_data = [
            ("课前", "课前自主学习", content.get("before_class", "")),
            ("课中", "导入", content.get("intro", "")),
            ("课中", "新课", content.get("teach", "")),
            ("课中", "小结", content.get("summary", "")),
            ("课后", "课后作业", content.get("homework", "")),
        ]
        fi = 0
        for ri in range(13, len(rows)):
            if fi >= len(flow_data):
                break
            row = rows[ri]
            phase, step, desc = flow_data[fi]
            if len(row.cells) >= 3:
                cell_content = row.cells[2]
                if not cell_content.text.strip() or len(cell_content.text.strip()) < 5:
                    _fill_cell(cell_content, desc[:150], 9)
            fi += 1

    url = _save_docx(doc, course_id, "教案")
    return {"success": True, "download_url": url}


# ========== 授课计划生成（整门课）==========

def generate_teaching_plan_docx(course_id):
    """基于授课计划.docx模板，生成整门课的授课计划"""
    from docx import Document
    from docx.shared import Pt

    tmpl = _template_path("授课计划.docx")
    if not tmpl:
        return {"success": False, "error": "授课计划模板不存在"}

    doc = Document(tmpl)
    cname = _course_name(course_id)

    # AI 生成（带降级）
    content = _ai_generate_teaching_plan(course_id)
    if not content:
        content = _fallback_teaching_plan(course_id)

    schedule = content.get("schedule", [])

    tables = doc.tables

    # Table0: 学时分配表（3行×9列）
    if len(tables) > 0:
        t0 = tables[0]
        rows = t0.rows
        if len(rows) > 2:
            data_row = rows[2]
            cells = data_row.cells
            if len(cells) >= 5:
                _fill_cell(cells[1], content.get("total_hours", "72"), 9)
                _fill_cell(cells[2], content.get("total_hours", "72"), 9)
                _fill_cell(cells[3], content.get("theory_hours", "24"), 9)
                _fill_cell(cells[4], content.get("practice_hours", "48"), 9)

    # Table1: 授课进度计划表（20行×11列）
    if len(tables) > 1:
        t1 = tables[1]
        for ri, row in enumerate(t1.rows):
            if ri < 2:
                continue
            idx = ri - 2
            if idx < len(schedule):
                s = schedule[idx]
                cells = row.cells
                if len(cells) >= 11:
                    _fill_cell(cells[0], str(s.get("week", idx+1)), 9)
                    _fill_cell(cells[1], s.get("class_num", f"{idx*4+1}-{(idx+1)*4}"), 9)
                    _fill_cell(cells[2], s.get("module", ""), 9)
                    _fill_cell(cells[3], s.get("task", ""), 9)
                    _fill_cell(cells[4], s.get("content", "")[:120], 9)
                    _fill_cell(cells[5], s.get("requirement", "")[:80], 9)
                    _fill_cell(cells[6], s.get("ideology", ""), 9)
                    _fill_cell(cells[7], str(s.get("theory_hours", 2)), 9)
                    _fill_cell(cells[8], str(s.get("practice_hours", 2)), 9)
                    _fill_cell(cells[9], s.get("method", "讲授法、讨论法、直观演示法、练习法"), 9)
                    _fill_cell(cells[10], "自拟", 9)

    url = _save_docx(doc, course_id, "授课计划")
    return {"success": True, "download_url": url}


# ========== 课程标准生成（整门课）==========

def generate_curriculum_standard_docx(course_id):
    """基于课标.docx模板，生成整门课的课程标准"""
    from docx import Document
    from docx.shared import Pt

    tmpl = _template_path("课标.docx")
    if not tmpl:
        return {"success": False, "error": "课标模板不存在"}

    doc = Document(tmpl)
    cname = _course_name(course_id)
    chapters = get_course_chapters(course_id) or []

    # AI 生成（带降级）
    content = _ai_generate_curriculum_standard(course_id)
    if not content:
        content = _fallback_curriculum_standard(course_id)

    # ===== 段落替换 =====
    for p in doc.paragraphs:
        for run in p.runs:
            if "程序设计基础（Python）" in run.text:
                run.text = run.text.replace("程序设计基础（Python）", cname)
            elif "《程序设计基础（Python）》" in run.text:
                run.text = run.text.replace("《程序设计基础（Python）》", f"《{cname}》")
            elif "物联网设备装调与维护" in run.text:
                run.text = run.text.replace("物联网设备装调与维护", cname)

    # ===== 表格替换 =====
    tables = doc.tables

    # Table0: 课程基本信息
    if len(tables) > 0:
        t0 = tables[0]
        for row in t0.rows:
            txt = " ".join(c.text.strip() for c in row.cells)
            if "课程代码" in txt:
                _fill_cell_by_keyword(row, "课程代码", "F61303", 9)
            elif "学分" in txt:
                _fill_cell_by_keyword(row, "学分", str(content.get("credits", 4)), 9)

    # Table2: 课程内容表
    if len(tables) > 2:
        t2 = tables[2]
        for ri, row in enumerate(t2.rows):
            if ri < 2:
                continue
            idx = ri - 2
            if idx < len(chapters):
                ch = chapters[idx]
                cells = row.cells
                if len(cells) >= 2:
                    _fill_cell(cells[1], ch.get('module', ''), 9)
                if len(cells) >= 3:
                    _fill_cell(cells[2], ch['title'], 9)

    # Table4: 教学设计子表
    if len(tables) > 4:
        t4 = tables[4]
        pi = 0
        for ri, row in enumerate(t4.rows):
            if ri == 0:
                continue
            txt = row.cells[0].text.strip() if row.cells else ""
            if not txt or "项目" in txt[:6]:
                continue
            if pi < len(chapters):
                ch = chapters[pi]
                if len(row.cells) >= 2 and not row.cells[1].text.strip():
                    _fill_cell(row.cells[1], ch['title'], 9)
                pi += 1

    url = _save_docx(doc, course_id, "课程标准")
    return {"success": True, "download_url": url}


# ========== 整门课教案生成 ==========

def generate_course_lesson_plan_docx(course_id):
    """生成整门课教案（所有章节合并到一个文档）"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tmpl = _template_path("教案.docx")
    if not tmpl:
        return {"success": False, "error": "教案模板不存在"}

    cname = _course_name(course_id)
    textbook = _course_textbook(course_id)
    chapters = get_course_chapters(course_id)
    if not chapters:
        return {"success": False, "error": "课程无章节"}

    # AI 生成（带降级）
    all_content = _ai_generate_course_lesson_plan(course_id)
    if not all_content:
        all_content = _fallback_course_lesson_plan(course_id)

    doc = Document(tmpl)
    for p in doc.paragraphs:
        p.clear()
    for table in list(doc.tables):
        table._element.getparent().remove(table._element)

    # ===== 封面 =====
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(f"{cname}  教  案")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run(f"（整门课 · 共 {len(chapters)} 章）")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_info.add_run(f"\n\n宜宾工业职业技术学院 · 数字经济学院\n授课教师：朱景峰")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    # ===== 逐章生成教案 =====
    for ci, ch in enumerate(chapters):
        content = None
        if all_content and ci < len(all_content):
            content = all_content[ci]

        h = doc.add_heading(f"第{ci+1}章  {ch['title']}", level=1)
        for run in h.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        # 基本信息表
        table = doc.add_table(rows=4, cols=4)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        info_data = [
            ["课程名称", cname, "授课主题", ch['title']],
            ["授课类型", "理实一体", "授课学时", "4学时"],
            ["教学内容分析", content.get("teaching_content_analysis", "见课程章节") if content else "见课程章节", "", ""],
            ["学情分析", content.get("student_analysis", "") if content else "", "", ""],
        ]
        for ri, row_data in enumerate(info_data):
            for ci2, val in enumerate(row_data):
                cell = table.rows[ri].cells[ci2]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                if ci2 == 0 or ci2 == 2:
                    run.bold = True

        doc.add_paragraph()

        # 教学目标表
        t2 = doc.add_table(rows=4, cols=2)
        try:
            t2.style = 'Table Grid'
        except Exception:
            pass

        fallback_ct = content if content else {}
        objectives_data = [
            ("知识目标", fallback_ct.get("knowledge_objective", ""),),
            ("能力目标", fallback_ct.get("ability_objective", ""),),
            ("素养目标", fallback_ct.get("literacy_objective", ""),),
            ("教学策略", fallback_ct.get("teaching_strategy", ""),),
        ]
        for ri, (label, val) in enumerate(objectives_data):
            cell0 = t2.rows[ri].cells[0]
            cell0.text = ""
            p = cell0.paragraphs[0]
            run = p.add_run(label)
            run.font.size = Pt(9)
            run.bold = True
            cell1 = t2.rows[ri].cells[1]
            cell1.text = ""
            p = cell1.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)

        doc.add_paragraph()

        # 教学活动表
        t3 = doc.add_table(rows=5, cols=3)
        try:
            t3.style = 'Table Grid'
        except Exception:
            pass

        headers = ["教学环节", "教学内容", "时间"]
        for ci3, h_text in enumerate(headers):
            cell = t3.rows[0].cells[ci3]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h_text)
            run.bold = True
            run.font.size = Pt(9)

        activities = [
            ("课前预习", fallback_ct.get("before_class", "")),
            ("导入", fallback_ct.get("intro", "")),
            ("讲授", fallback_ct.get("teach", "")),
            ("巩固练习", fallback_ct.get("practice", "")),
        ]
        for ri, (phase, desc) in enumerate(activities):
            row = t3.rows[ri + 1]
            row.cells[0].text = ""
            p = row.cells[0].paragraphs[0]
            run = p.add_run(phase)
            run.font.size = Pt(9)
            run.bold = True
            row.cells[1].text = ""
            p = row.cells[1].paragraphs[0]
            run = p.add_run(desc[:100])
            run.font.size = Pt(9)
            row.cells[2].text = ""
            p = row.cells[2].paragraphs[0]
            run = p.add_run("4学时" if ri < 2 else "")

        if ci < len(chapters) - 1:
            doc.add_page_break()

    url = _save_docx(doc, course_id, "整门课教案")
    return {"success": True, "download_url": url, "chapter_count": len(chapters)}
